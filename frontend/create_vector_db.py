import os
import streamlit as st
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import DirectoryLoader, PyPDFLoader, Docx2txtLoader
import docx
from PyPDF2 import PdfReader, PdfWriter
from io import StringIO
import pdfplumber



VECTOR_DB_PATH = "faiss_index"
DATA_PATH = "data"  # Path to the directory containing PDF and DOCX files

if 'GOOGLE_API_KEY' not in os.environ:
    os.environ["GOOGLE_API_KEY"] = "AIzaSyA-MiyxXeCTQ0Nb8j86g67S65vTUaimZLk"


def list_files():
    return [f for f in os.listdir(DATA_PATH) if os.path.isfile(os.path.join(DATA_PATH, f))]


def delete_file(file_name):
    file_path = os.path.join(DATA_PATH, file_name)
    if os.path.exists(file_path):
        os.remove(file_path)


def load_docx(file_path):
    doc = docx.Document(file_path)
    return '\n'.join([para.text for para in doc.paragraphs])


def save_docx(file_path, content):
    doc = docx.Document()
    for line in content.split('\n'):
        doc.add_paragraph(line)
    doc.save(file_path)


def load_pdf(file_path):
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    return text


search_results = []  # Biến lưu trữ kết quả tìm kiếm
def search_text(content, search_term):
    return content.replace(search_term, f"**{search_term}**")

def save_pdf(file_path, content):
    writer = PdfWriter()
    buffer = StringIO(content)
    reader = PdfReader(buffer)
    for page in reader.pages:
        writer.add_page(page)
    with open(file_path, "wb") as output_file:
        writer.write(output_file)


def main():
    st.title("Quản Lý Cơ Sở Dữ Liệu Vector FAISS")

    st.header("Quản lý tệp")

    # List files
    files = list_files()
    st.subheader("Danh sách tệp hiện có")
    selected_file = st.selectbox("Chọn tệp", files)

    # Upload files
    st.subheader("Thêm tệp mới")
    uploaded_files = st.file_uploader("Chọn tệp PDF hoặc DOCX", accept_multiple_files=True, type=['pdf', 'docx'])
    if uploaded_files:
        existing_files = list_files()
        duplicate_files = [file.name for file in uploaded_files if file.name in existing_files]

        if duplicate_files:
            st.error(f"Các tệp sau đã tồn tại: {', '.join(duplicate_files)}")
        else:
            for uploaded_file in uploaded_files:
                if st.button(f"Cập nhật: {uploaded_file.name}"):
                    with open(os.path.join(DATA_PATH, uploaded_file.name), 'wb') as f:
                        f.write(uploaded_file.getbuffer())
                        st.success(f"Tệp {uploaded_file.name} đã được tải lên thành công.")
        
    # Delete files
    st.subheader("Xóa tệp")
    file_to_delete = st.selectbox("Chọn tệp để xóa", files)
    if st.button("Xóa tệp"):
        delete_file(file_to_delete)
        st.success(f"Tệp {file_to_delete} đã được xóa.")
        files = list_files()
        file_path = os.path.join(DATA_PATH, selected_file)
        file_extension = os.path.splitext(selected_file)[1].lower()
    st.header("Cập nhật cơ sở dữ liệu vector FAISS")
    if st.button("Cập nhật Cơ Sở Dữ Liệu Vector"):
        # Initialize loaders
        pdf_loader = DirectoryLoader(DATA_PATH, glob="*.pdf", loader_cls=PyPDFLoader)
        docx_loader = DirectoryLoader(DATA_PATH, glob="*.docx", loader_cls=Docx2txtLoader)

        # Load PDF and DOCX documents
        pdf_files = pdf_loader.load()
        docx_files = docx_loader.load()

        # Concatenate text from PDF and DOCX files
        text = ""
        for pdf_doc in pdf_files:
            text += pdf_doc.page_content
        for docx_doc in docx_files:
            text += docx_doc.page_content

        # Split text into chunks
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=7000, chunk_overlap=1400)
        chunks = text_splitter.split_text(text)

        # Initialize embeddings
        embeddings = GoogleGenerativeAIEmbeddings(model="models/text-embedding-004")

        # Create FAISS database
        db = FAISS.from_texts(chunks, embeddings)
        db.save_local(VECTOR_DB_PATH)

        st.success("Cơ sở dữ liệu vector đã được cập nhật thành công.")

if __name__ == "__main__":
    main()
